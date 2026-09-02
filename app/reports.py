from __future__ import annotations

import csv, json, os
from pathlib import Path
from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether

ROOT=Path(__file__).resolve().parents[1]
DATA_DIR=Path(os.getenv('DATA_DIR') or os.getenv('RAILWAY_VOLUME_MOUNT_PATH') or (ROOT/'data'))
OUT=DATA_DIR/'exports'
OUT.mkdir(parents=True,exist_ok=True)


def _register_fonts() -> tuple[str,str]:
    regular_candidates=[Path(r'C:\Windows\Fonts\msyh.ttc'),Path('/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'),Path('/System/Library/Fonts/PingFang.ttc')]
    bold_candidates=[Path(r'C:\Windows\Fonts\msyhbd.ttc'),Path('/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc')]
    regular='Helvetica';bold='Helvetica-Bold'
    for p in regular_candidates:
        if p.exists():
            try:pdfmetrics.registerFont(TTFont('InsightCJK',str(p),subfontIndex=0));regular='InsightCJK';break
            except Exception:
                try:pdfmetrics.registerFont(TTFont('InsightCJK',str(p)));regular='InsightCJK';break
                except Exception:pass
    for p in bold_candidates:
        if p.exists():
            try:pdfmetrics.registerFont(TTFont('InsightCJKBold',str(p),subfontIndex=0));bold='InsightCJKBold';break
            except Exception:
                try:pdfmetrics.registerFont(TTFont('InsightCJKBold',str(p)));bold='InsightCJKBold';break
                except Exception:pass
    return regular,bold


def _clean(s) -> str:
    return str(s if s is not None else '—').replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')


def _styles():
    regular,bold=_register_fonts();base=getSampleStyleSheet()
    styles={
        'title':ParagraphStyle('title',parent=base['Title'],fontName=bold,fontSize=22,leading=27,textColor=colors.HexColor('#11161d'),spaceAfter=7*mm,alignment=TA_LEFT),
        'kicker':ParagraphStyle('kicker',parent=base['Normal'],fontName=bold,fontSize=8,leading=10,textColor=colors.HexColor('#147a63'),spaceAfter=2*mm),
        'h1':ParagraphStyle('h1',parent=base['Heading1'],fontName=bold,fontSize=15,leading=19,textColor=colors.HexColor('#11161d'),spaceBefore=3*mm,spaceAfter=3*mm),
        'h2':ParagraphStyle('h2',parent=base['Heading2'],fontName=bold,fontSize=11,leading=14,textColor=colors.HexColor('#172027'),spaceBefore=2*mm,spaceAfter=2*mm),
        'body':ParagraphStyle('body',parent=base['BodyText'],fontName=regular,fontSize=8.7,leading=13,textColor=colors.HexColor('#4f5964'),spaceAfter=2.3*mm),
        'small':ParagraphStyle('small',parent=base['BodyText'],fontName=regular,fontSize=7.3,leading=10,textColor=colors.HexColor('#7b8690')),
        'strong':ParagraphStyle('strong',parent=base['BodyText'],fontName=bold,fontSize=9.2,leading=13,textColor=colors.HexColor('#1f2933')),
    }
    return styles,regular,bold


def _footer(canvas,doc):
    canvas.saveState();regular,_=_register_fonts();canvas.setFont(regular,7);canvas.setFillColor(colors.HexColor('#8d969f'))
    canvas.drawString(18*mm,10*mm,'InsightFlow AI · Evidence-backed Consumer & GTM Intelligence')
    canvas.drawRightString(A4[0]-18*mm,10*mm,f'Page {doc.page}');canvas.restoreState()


def write_docx(research: dict, summary: dict) -> Path:
    # Backwards-compatible export. Recruiter UI intentionally prioritizes PDF + evidence CSV.
    p=OUT/f"insightflow_research_{research['id']}.docx";doc=Document();doc.add_heading('InsightFlow AI · Consumer & GTM Intelligence Brief',0)
    doc.add_paragraph(f"Research: {research['keyword']}")
    doc.add_paragraph(f"Markets: {research['markets_json']} · Window: {research['days']} days")
    demo=((summary.get('research') or {}).get('decision') or {}).get('demo_meta') or {}
    if demo.get('executive_recommendation'):
        doc.add_heading('Decision',1);doc.add_paragraph(demo['executive_recommendation'])
    c=summary.get('evidence_confidence') or {};doc.add_heading('Evidence quality',1);doc.add_paragraph(f"Confidence: {c.get('label','—')} ({c.get('score','—')}/100). Consumer evidence: {summary['review_count']}; product signals: {summary['product_count']}.")
    doc.add_heading('Top consumer topics',1)
    for x in summary.get('issues',[])[:6]:doc.add_paragraph(f"{x['name']}: {x['count']} rows; sample share {x['share']}%; decision-impact {x['purchase_impact_rate']}%; {x['confidence']} confidence",style='List Bullet')
    doc.add_heading('Opportunity hypotheses',1)
    for x in summary.get('opportunities',[])[:5]:
        d=x.get('decision') or {};doc.add_paragraph(f"{x['name']} — priority {x['opportunity_score']}/100; {x['confidence']} confidence",style='List Bullet')
        if d.get('product_action'):doc.add_paragraph('Product: '+d['product_action'])
        if d.get('gtm_action'):doc.add_paragraph('GTM: '+d['gtm_action'])
        if d.get('next_validation'):doc.add_paragraph('Next validation: '+d['next_validation'])
    doc.add_heading('Method boundary',1);doc.add_paragraph('No synthetic-review fallback. Review count is not sales. GLOBAL consumer evidence is not assigned to a country. Cross-market claims are blocked when equivalent consumer-source coverage is absent. Opportunity priority is not TAM or revenue estimation.')
    doc.save(p);return p


def write_pdf(research: dict, summary: dict) -> Path:
    p=OUT/f"insightflow_research_{research['id']}.pdf";styles,regular,bold=_styles()
    doc=SimpleDocTemplate(str(p),pagesize=A4,rightMargin=17*mm,leftMargin=17*mm,topMargin=18*mm,bottomMargin=17*mm,title='InsightFlow AI Consumer & GTM Intelligence Brief')
    story=[];rr=summary.get('research') or {};decision=(rr.get('decision') or {});demo=decision.get('demo_meta') or {};conf=summary.get('evidence_confidence') or {}
    story += [Paragraph('EXECUTIVE RESEARCH BRIEF',styles['kicker']),Paragraph('InsightFlow AI · Consumer & GTM Intelligence',styles['title'])]
    meta=[[Paragraph('<b>Research</b><br/>'+_clean(research.get('keyword')),styles['body']),Paragraph('<b>Markets</b><br/>'+_clean(', '.join(rr.get('markets') or [])),styles['body']),Paragraph('<b>Evidence confidence</b><br/>'+_clean(f"{conf.get('label','—')} · {conf.get('score','—')}/100"),styles['body'])]]
    t=Table(meta,colWidths=[72*mm,45*mm,47*mm]);t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),colors.HexColor('#f5f7f8')),('BOX',(0,0),(-1,-1),.5,colors.HexColor('#dce2e6')),('INNERGRID',(0,0),(-1,-1),.25,colors.HexColor('#e7ebee')),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),8),('RIGHTPADDING',(0,0),(-1,-1),8),('TOPPADDING',(0,0),(-1,-1),8),('BOTTOMPADDING',(0,0),(-1,-1),8)]));story += [t,Spacer(1,6*mm)]
    if demo.get('business_question'):
        story += [Paragraph('BUSINESS QUESTION',styles['kicker']),Paragraph(_clean(demo['business_question']),styles['strong']),Spacer(1,3*mm)]
    if demo.get('executive_recommendation'):
        story += [Paragraph('DECISION',styles['kicker']),Paragraph(_clean(demo['executive_recommendation']),styles['h1'])]
    story += [Paragraph(f"Current evidence: {summary.get('review_count',0)} consumer rows · {summary.get('product_count',0)} product signals · {len(summary.get('sources') or {})} consumer source(s).",styles['small']),Spacer(1,4*mm)]

    story += [Paragraph('What the evidence says',styles['h1'])]
    issue_rows=[[Paragraph('Topic',styles['small']),Paragraph('Evidence',styles['small']),Paragraph('Sample share',styles['small']),Paragraph('Decision impact',styles['small']),Paragraph('Confidence',styles['small'])]]
    for x in (summary.get('issues') or [])[:6]:issue_rows.append([Paragraph(_clean(x['name']),styles['body']),str(x['count']),f"{x['share']}%",f"{x['purchase_impact_rate']}%",_clean(x['confidence'])])
    it=Table(issue_rows,colWidths=[58*mm,20*mm,28*mm,30*mm,25*mm],repeatRows=1);it.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#152028')),('TEXTCOLOR',(0,0),(-1,0),colors.white),('FONTNAME',(0,0),(-1,0),bold),('FONTNAME',(0,1),(-1,-1),regular),('FONTSIZE',(0,0),(-1,-1),7.5),('GRID',(0,0),(-1,-1),.3,colors.HexColor('#dde3e7')),('VALIGN',(0,0),(-1,-1),'TOP'),('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,colors.HexColor('#f8fafb')]),('LEFTPADDING',(0,0),(-1,-1),6),('RIGHTPADDING',(0,0),(-1,-1),6),('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6)]));story += [it,Spacer(1,5*mm)]

    # Recruiter brief should show only action cards that have a real decision memo attached.
    # Generic scored topics remain visible in the evidence table, but an empty action card
    # would imply a recommendation the system did not actually make.
    scored_by_name={x.get('name'):x for x in (summary.get('opportunities') or []) if x.get('name')}
    decision_actions=[]
    for d in (decision.get('opportunities') or []):
        name=d.get('name')
        if not name:
            continue
        x=dict(scored_by_name.get(name) or {'name':name,'opportunity_score':'—','confidence':'Directional','benchmark_coverage':'—'})
        x['decision']=d
        decision_actions.append(x)

    story += [PageBreak(),Paragraph('Priority actions',styles['h1'])]
    for i,x in enumerate(decision_actions[:3],1):
        d=x.get('decision') or {};blocks=[Paragraph(f"0{i} · {_clean(x.get('name'))}",styles['h2']),Paragraph(f"Priority {x.get('opportunity_score','—')}/100 · {_clean(x.get('confidence'))} confidence · benchmark coverage proxy {x.get('benchmark_coverage','—')}%",styles['small'])]
        if d.get('insight'):blocks.append(Paragraph('<b>Insight:</b> '+_clean(d['insight']),styles['body']))
        if d.get('product_action'):blocks.append(Paragraph('<b>Product action:</b> '+_clean(d['product_action']),styles['body']))
        if d.get('gtm_action'):blocks.append(Paragraph('<b>GTM action:</b> '+_clean(d['gtm_action']),styles['body']))
        if d.get('next_validation'):blocks.append(Paragraph('<b>Next validation:</b> '+_clean(d['next_validation']),styles['body']))
        box=Table([[blocks]],colWidths=[164*mm]);box.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),colors.HexColor('#f7faf9')),('BOX',(0,0),(-1,-1),.6,colors.HexColor('#d9e8e2')),('LEFTPADDING',(0,0),(-1,-1),10),('RIGHTPADDING',(0,0),(-1,-1),10),('TOPPADDING',(0,0),(-1,-1),8),('BOTTOMPADDING',(0,0),(-1,-1),8)]));story += [KeepTogether(box),Spacer(1,3*mm)]

    story += [PageBreak(),Paragraph('Evidence quality & decision boundaries',styles['h1'])]
    mc=summary.get('market_comparison') or {};coverage=summary.get('source_coverage') or {}
    story += [Paragraph('<b>Cross-market consumer comparison:</b> '+('Available from '+_clean(', '.join(mc.get('common_sources') or [])) if mc.get('available') else 'Blocked - '+_clean(mc.get('reason') or 'coverage gap')),styles['body'])]
    story += [Paragraph('<b>Trend interpretation:</b> Google Trends values are normalized within each market. Separately requested market indices are not treated as absolute market-size ratios.',styles['body'])]
    story += [Paragraph('<b>Sales boundary:</b> Product review counts and ranking signals are not presented as unit sales.',styles['body'])]
    story += [Paragraph('<b>AI boundary:</b> LLM is used to structure supplied evidence and draft actions. Missing source evidence is never replaced by synthetic consumer rows.',styles['body']),Spacer(1,3*mm)]
    if coverage.get('rows'):
        story += [Paragraph('Source coverage',styles['h2'])]
        cov_rows=[[Paragraph('Source',styles['small']),Paragraph('Evidence type',styles['small']),Paragraph('Coverage',styles['small']),Paragraph('Cross-market',styles['small'])]]
        for r in coverage['rows']:
            counts=' · '.join(f"{k}:{v}" for k,v in (r.get('counts') or {}).items())
            cov_rows.append([_clean(r.get('source')),_clean(r.get('kind')), _clean(counts), 'Comparable' if r.get('comparable_across_requested_markets') else 'Gap'])
        ct=Table(cov_rows,colWidths=[48*mm,35*mm,50*mm,30*mm],repeatRows=1);ct.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#eef2f4')),('FONTNAME',(0,0),(-1,0),bold),('FONTNAME',(0,1),(-1,-1),regular),('FONTSIZE',(0,0),(-1,-1),7.2),('GRID',(0,0),(-1,-1),.3,colors.HexColor('#dfe4e8')),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)]));story += [ct]
    if demo.get('limitations'):
        story += [Spacer(1,5*mm),Paragraph('Known limitations',styles['h2'])]
        for x in demo['limitations']:story.append(Paragraph('• '+_clean(x),styles['body']))
    story += [Spacer(1,4*mm),Paragraph('Interpretation rule',styles['h2']),Paragraph(_clean(demo.get('decision_principle') or 'Make the strongest decision the evidence supports - and expose what still needs validation.'),styles['strong'])]
    doc.build(story,onFirstPage=_footer,onLaterPages=_footer);return p


def write_csv(reviews: list[dict], research_id: int) -> Path:
    p=OUT/f'insightflow_evidence_{research_id}.csv'
    fields=['source','market','product_title','title','text','rating','author','review_date','url','helpful','sentiment','issue','topics_json','driver','barrier','purchase_impact','scenario','competitor_mentions_json','analysis_mode','window_status']
    def safe_cell(v):
        if isinstance(v,str) and v.lstrip().startswith(('=','+','-','@')):return "'"+v
        return v
    with p.open('w',encoding='utf-8-sig',newline='') as f:
        w=csv.DictWriter(f,fieldnames=fields,extrasaction='ignore');w.writeheader()
        for row in reviews:w.writerow({k:safe_cell(row.get(k)) for k in fields})
    return p
